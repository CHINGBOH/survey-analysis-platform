"""app/reports.py — 报告生成系统 (Word / PDF / 图片包 / 模板)

支持:
  - Quarto HTML (默认,已有 04-report/report.qmd)
  - Word (python-docx + jinja-like 模板)
  - PDF (Quarto + xelatex 中文,或 reportlab 后备)
  - 图片包 (zip,所有 PNG+SVG+CSV)
  - 多模板 (简版/标准版/详尽版)
"""
from __future__ import annotations
import os, json, zipfile, shutil, subprocess, glob
from pathlib import Path
from datetime import datetime

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output"
REPORTS = OUTPUT / "reports"
TEMPLATES = ROOT / "04-report" / "templates"

REPORTS.mkdir(parents=True, exist_ok=True)
TEMPLATES.mkdir(parents=True, exist_ok=True)

MODULE_TITLES = {
    "descriptives": "描述统计", "crosstabs": "交叉表分析",
    "ttest": "t 检验", "anova": "方差分析", "correlation": "相关分析",
    "reliability": "信度分析", "factor_analysis": "因子分析",
    "regression": "回归分析", "mediation": "中介效应",
    "moderation": "调节效应", "cluster": "聚类分析",
    "power_bootstrap": "Bootstrap 与效力", "survey_specific": "问卷专用分析",
}

TEMPLATE_SPECS = {
    "minimal": {"name": "简版", "modules": ["descriptives", "crosstabs"],
                "include_charts": False, "include_appendix": False},
    "standard": {"name": "标准版",
                "modules": ["descriptives", "crosstabs", "ttest", "anova",
                            "correlation", "regression", "reliability"],
                "include_charts": True, "include_appendix": False},
    "full": {"name": "详尽版",
            "modules": list(MODULE_TITLES.keys()),
            "include_charts": True, "include_appendix": True},
}


# ─── Word 报告 ──────────────────────────────────────────────
def generate_word_report(
    template: str = "standard",
    org_name: str = "调查分析平台",
    period: str = "",
    title: str = "问卷调查分析报告",
    survey_label: str = "s1",
    out_path: str | None = None,
) -> dict:
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return {"ok": False, "error": "python-docx 未安装,运行: pip install python-docx"}

    spec = TEMPLATE_SPECS.get(template, TEMPLATE_SPECS["standard"])
    period = period or datetime.now().strftime("%Y-%m")
    out_path = out_path or str(REPORTS / f"report_{template}_{survey_label}.docx")

    doc = Document()

    # 封面
    title_p = doc.add_heading(title, level=0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"\n{org_name}\n").bold = True
    p.add_run(f"\n报告期: {period}\n报告类型: {spec['name']}\n")
    p.add_run(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_page_break()

    # 目录(简易 — 直接列模块)
    doc.add_heading("目录", level=1)
    for i, m in enumerate(spec["modules"], 1):
        doc.add_paragraph(f"{i}. {MODULE_TITLES.get(m, m)}", style="List Number")
    doc.add_page_break()

    # 各模块
    for i, mod in enumerate(spec["modules"], 1):
        doc.add_heading(f"{i}. {MODULE_TITLES.get(mod, mod)}", level=1)
        # 数据摘要 (从 RDS 转出的 JSON,如不存在则跳过)
        json_path = OUTPUT / "results" / f"{mod}_{survey_label}.json"
        if json_path.exists():
            try:
                with open(json_path, encoding="utf-8") as f:
                    data = json.load(f)
                _add_module_to_doc(doc, mod, data, with_charts=spec["include_charts"],
                                   survey_label=survey_label)
            except Exception as e:
                doc.add_paragraph(f"[无法解析结果: {e}]")
        else:
            doc.add_paragraph(f"(本模块尚无结果文件: {json_path.name})", style="Intense Quote")

        # 图表
        if spec["include_charts"]:
            chart_dir = OUTPUT / "charts" / f"{mod}_{survey_label}"
            if chart_dir.exists():
                pngs = sorted(chart_dir.glob("*.png"))
                for png in pngs[:6]:
                    try:
                        doc.add_picture(str(png), width=Inches(5.5))
                        cap = doc.add_paragraph(png.stem.replace("_", " "))
                        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cap.runs[0].italic = True
                    except Exception:
                        pass

    # 附录
    if spec["include_appendix"]:
        doc.add_page_break()
        doc.add_heading("附录: 技术说明", level=1)
        doc.add_paragraph(
            "本报告由 survey-analysis-platform 自动生成。所有统计方法对标 SPSS 标准:\n"
            "  · t 检验: Levene + Student + Welch + Cohen's d\n"
            "  · ANOVA: Welch + Tukey HSD + Games-Howell + Kruskal-Wallis\n"
            "  · 回归: VIF + DW + BP + Cook + Hosmer-Lemeshow\n"
            "  · 因子: KMO + Bartlett + Varimax 旋转\n"
            "  · 信度: Cronbach α + Split-half + 项目分析"
        )

    doc.save(out_path)
    return {"ok": True, "path": out_path, "template": template, "modules": len(spec["modules"])}


def _add_module_to_doc(doc, module, data, with_charts=True, survey_label="s1"):
    """把模块 JSON 结果按主表渲染为 Word 表格"""
    from docx.shared import Pt
    # 找出主表 (data.frame 序列化为列表[dict])
    rendered = False
    for key, val in (data or {}).items():
        if isinstance(val, list) and val and isinstance(val[0], dict):
            doc.add_paragraph(key, style="Heading 3")
            _df_to_table(doc, val)
            rendered = True
        elif isinstance(val, dict):
            for k2, v2 in val.items():
                if isinstance(v2, list) and v2 and isinstance(v2[0], dict):
                    doc.add_paragraph(f"{key} / {k2}", style="Heading 3")
                    _df_to_table(doc, v2)
                    rendered = True
    if not rendered:
        doc.add_paragraph("(结果为非表格结构,详见 HTML 报告)", style="Intense Quote")


def _df_to_table(doc, rows, max_rows=30):
    """list[dict] → Word table"""
    if not rows: return
    cols = list(rows[0].keys())
    tbl = doc.add_table(rows=1, cols=len(cols))
    tbl.style = "Light Grid"
    hdr = tbl.rows[0].cells
    for i, c in enumerate(cols): hdr[i].text = str(c)
    for r in rows[:max_rows]:
        row = tbl.add_row().cells
        for i, c in enumerate(cols):
            v = r.get(c, "")
            row[i].text = "" if v is None else str(v)


# ─── PDF (尝试 Quarto/LaTeX,失败退化为 reportlab 原生 PDF) ──
def generate_pdf_report(template: str = "standard", survey_label: str = "s1",
                        title: str = "问卷调查分析报告", org_name: str = "调查分析平台") -> dict:
    """PDF: 优先 Quarto+xelatex; 退化为 reportlab 原生(中文 Noto)"""
    out_path = REPORTS / f"report_{template}_{survey_label}.pdf"

    # 尝试 quarto+latex
    qmd = ROOT / "04-report" / "report.qmd"
    if shutil.which("xelatex") and qmd.exists():
        try:
            r = subprocess.run(
                ["quarto", "render", str(qmd), "--to", "pdf",
                 "--output", out_path.name, "--output-dir", str(REPORTS)],
                capture_output=True, text=True, timeout=180, cwd=ROOT
            )
            if r.returncode == 0 and out_path.exists():
                return {"ok": True, "path": str(out_path), "engine": "quarto+xelatex"}
        except Exception:
            pass

    # 退化: reportlab 原生
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Image,
                                         Table, TableStyle, PageBreak)
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        # 中文字体
        font_name = "Helvetica"
        for f in ["/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
                  "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
                  "/usr/share/fonts/opentype/noto/NotoSerifCJK-Regular.ttc"]:
            if os.path.exists(f):
                try:
                    pdfmetrics.registerFont(TTFont("NotoCJK", f))
                    font_name = "NotoCJK"
                    break
                except Exception:
                    continue

        spec = TEMPLATE_SPECS.get(template, TEMPLATE_SPECS["standard"])
        doc = SimpleDocTemplate(str(out_path), pagesize=A4,
                                leftMargin=2*cm, rightMargin=2*cm,
                                topMargin=2*cm, bottomMargin=2*cm)
        styles = getSampleStyleSheet()
        h0 = ParagraphStyle("H0", parent=styles["Title"], fontName=font_name, fontSize=22, alignment=1)
        h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontName=font_name, fontSize=16)
        h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontName=font_name, fontSize=13)
        body = ParagraphStyle("Body", parent=styles["Normal"], fontName=font_name, fontSize=10, leading=14)
        story = []

        story += [Spacer(1, 4*cm), Paragraph(title, h0), Spacer(1, 1*cm),
                  Paragraph(org_name, body), Spacer(1, 0.3*cm),
                  Paragraph(f"报告类型: {spec['name']}", body),
                  Paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}", body),
                  PageBreak()]

        story += [Paragraph("目录", h1)]
        for i, m in enumerate(spec["modules"], 1):
            story.append(Paragraph(f"{i}. {MODULE_TITLES.get(m, m)}", body))
        story.append(PageBreak())

        for i, mod in enumerate(spec["modules"], 1):
            story.append(Paragraph(f"{i}. {MODULE_TITLES.get(mod, mod)}", h1))
            json_path = OUTPUT / "results" / f"{mod}_{survey_label}.json"
            if json_path.exists():
                try:
                    data = json.loads(json_path.read_text(encoding="utf-8"))
                    _add_module_to_pdf(story, mod, data, font_name, h2, body)
                except Exception as e:
                    story.append(Paragraph(f"[解析失败: {e}]", body))
            if spec["include_charts"]:
                chart_dir = OUTPUT / "charts" / f"{mod}_{survey_label}"
                if chart_dir.exists():
                    for png in sorted(chart_dir.glob("*.png"))[:4]:
                        try:
                            story.append(Spacer(1, 0.3*cm))
                            story.append(Image(str(png), width=14*cm, height=9*cm))
                            cap = Paragraph(f"<i>{png.stem.replace('_', ' ')}</i>", body)
                            story.append(cap)
                        except Exception:
                            pass
            story.append(Spacer(1, 0.5*cm))

        doc.build(story)
        return {"ok": True, "path": str(out_path), "engine": "reportlab"}
    except Exception as e:
        return {"ok": False, "error": f"PDF 生成失败: {e}"}


def _add_module_to_pdf(story, mod, data, font_name, h2, body):
    from reportlab.platypus import Paragraph, Table, TableStyle, Spacer
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    rendered = False
    for key, val in (data or {}).items():
        if isinstance(val, list) and val and isinstance(val[0], dict):
            story.append(Paragraph(str(key), h2))
            _list_to_pdf_table(story, val, font_name)
            rendered = True
        elif isinstance(val, dict):
            for k2, v2 in val.items():
                if isinstance(v2, list) and v2 and isinstance(v2[0], dict):
                    story.append(Paragraph(f"{key} / {k2}", h2))
                    _list_to_pdf_table(story, v2, font_name)
                    rendered = True
    if not rendered:
        story.append(Paragraph("(非表格结构,详见 HTML 报告)", body))


def _list_to_pdf_table(story, rows, font_name, max_rows=20):
    from reportlab.platypus import Table, TableStyle, Spacer
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    if not rows: return
    cols = list(rows[0].keys())
    data = [cols] + [[str(r.get(c, "")) for c in cols] for r in rows[:max_rows]]
    t = Table(data, repeatRows=1)
    t.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), font_name),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ("GRID", (0, 0), (-1, -1), 0.3, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    story.append(t)
    story.append(Spacer(1, 0.3*cm))


# ─── 图片包导出 ─────────────────────────────────────────────
def export_image_bundle(survey_label: str = "s1", formats=("png",), dpi=300) -> dict:
    """打包所有图表为 zip,可选导出多种格式 + CSV 数据"""
    bundle_dir = REPORTS / f"images_{survey_label}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    bundle_dir.mkdir(parents=True, exist_ok=True)

    charts_root = OUTPUT / "charts"
    n_images = 0
    chart_dirs = [d for d in charts_root.glob(f"*_{survey_label}") if d.is_dir()]
    for cd in chart_dirs:
        mod = cd.name.replace(f"_{survey_label}", "")
        sub = bundle_dir / mod
        sub.mkdir(exist_ok=True)
        for png in cd.glob("*.png"):
            shutil.copy(png, sub / png.name)
            n_images += 1
        # 复制 manifest
        mf = cd / "manifest.json"
        if mf.exists(): shutil.copy(mf, sub / "manifest.json")

    # 缩略图清单 PDF (简易: 用 PIL 拼成大图)
    try:
        from PIL import Image
        all_pngs = sorted(bundle_dir.rglob("*.png"))
        if all_pngs:
            thumbs = []
            for p in all_pngs:
                try:
                    img = Image.open(p).convert("RGB")
                    img.thumbnail((400, 300))
                    thumbs.append((img, p.parent.name + "/" + p.name))
                except Exception: pass
            if thumbs:
                # 拼图:每行 3 张
                tw, th = 420, 320
                cols = 3
                rows = (len(thumbs) + cols - 1) // cols
                canvas = Image.new("RGB", (tw * cols, th * rows), "white")
                for i, (im, _name) in enumerate(thumbs):
                    canvas.paste(im, ((i % cols) * tw, (i // cols) * th))
                canvas.save(bundle_dir / "_thumbnails.pdf", "PDF", resolution=100)
    except ImportError:
        pass

    # zip
    zip_path = REPORTS / f"images_{survey_label}.zip"
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for f in bundle_dir.rglob("*"):
            if f.is_file():
                zf.write(f, f.relative_to(bundle_dir))

    return {"ok": True, "zip": str(zip_path), "dir": str(bundle_dir),
            "n_images": n_images, "n_modules": len(chart_dirs)}


# ─── 模板系统 ────────────────────────────────────────────────
def list_templates() -> list[dict]:
    return [{"key": k, **v} for k, v in TEMPLATE_SPECS.items()]


def export_results_to_json(survey_label: str = "s1") -> dict:
    """遍历各模块 RDS 生成 JSON 供 Word/Python 渲染"""
    results_dir = OUTPUT / "results"
    converted = []
    for rds in results_dir.glob(f"*_{survey_label}.rds"):
        json_path = rds.with_suffix(".json")
        r = subprocess.run(
            ["Rscript", "02-analyze/rds_to_json.R", str(rds.relative_to(ROOT)), "30"],
            capture_output=True, text=True, cwd=ROOT, timeout=30
        )
        if r.returncode == 0 and r.stdout.strip():
            try:
                json_path.write_text(r.stdout, encoding="utf-8")
                converted.append(rds.stem)
            except Exception:
                pass
    return {"ok": True, "converted": converted, "n": len(converted)}
